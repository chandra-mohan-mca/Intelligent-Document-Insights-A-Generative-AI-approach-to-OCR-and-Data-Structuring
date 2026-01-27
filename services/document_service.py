from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io

def create_word_document(text: str) -> io.BytesIO:
    """
    Creates a Word document from tagged text using [TAM], [ENG], [PARA], [ADDRESS] tags.
    Implements a strict 'Ditto' match layout.
    """
    doc = Document()
    
    # Configure page settings
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.right_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Base style
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(12) # Gap between blocks
    style.paragraph_format.line_spacing = 1.15

    lines = text.split('\n')
    left_footer = []
    right_footer = []
    seen_para = False
    
    for line in lines:
        if not line.strip():
            continue
            
        # Parse tags
        current_tag = '[LEFT]' # Default
        if '[CENTER]' in line: current_tag = '[CENTER]'
        elif '[RIGHT]' in line: current_tag = '[RIGHT]'
        elif '[PARA]' in line: current_tag = '[PARA]'
        elif '[LEFT]' in line: current_tag = '[LEFT]'

        if current_tag == '[PARA]':
            seen_para = True

        is_tamil = '[TAM]' in line or any('\u0b80' <= char <= '\u0bff' for char in line)
        is_english = '[ENG]' in line
        
        # Clean text
        clean_text = line
        tags_to_remove = ['[LEFT]', '[CENTER]', '[RIGHT]', '[PARA]', '[TAM]', '[ENG]', '[ADDRESS]']
        for tag in tags_to_remove:
            clean_text = clean_text.replace(tag, '')
        clean_text = clean_text.strip()
        
        if not clean_text:
            continue

        # FOOTER DETECTION: Only trigger for footer elements after at least one paragraph or if tagged specifically
        # Signature block is always [RIGHT]
        if current_tag == '[RIGHT]':
            right_footer.append((clean_text, is_tamil, is_english))
            continue
        
        # Date/Place are [LEFT] but should only be footer if they appear after body text
        if seen_para and current_tag == '[LEFT]' and ('நாள்' in clean_text or 'இடம்' in clean_text):
            left_footer.append((clean_text, is_tamil, is_english))
            continue

        p = doc.add_paragraph()
        fmt = p.paragraph_format
        
        # Alignment & Indent
        if current_tag == '[CENTER]':
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.first_line_indent = Pt(0)
        elif current_tag == '[PARA]':
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = Inches(0.5)
        else:
            fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = Pt(0)
            
        # SPACING LOGIC: Tight blocks for everything except major paragraphs
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(12) if current_tag == '[PARA]' or current_tag == '[CENTER]' else Pt(0)
        fmt.line_spacing = 1.15
        
        # SMART FONT SWITCHING: Split text into Tamil and Non-Tamil segments
        # This prevents commas and numbers from turning into Tamil characters
        import re
        segments = re.split(r'([\u0b80-\u0bff]+)', clean_text)
        
        for segment in segments:
            if not segment:
                continue
            run = p.add_run(segment)
            run.font.size = Pt(12)
            
            # If segment is Tamil, apply SunTommy with full protection
            if any('\u0b80' <= char <= '\u0bff' for char in segment):
                run.font.name = "SunTommy"
                rPr = run.font._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:cs'), "SunTommy")
                rFonts.set(qn('w:ascii'), "SunTommy")
                rFonts.set(qn('w:hAnsi'), "SunTommy")
            else:
                # Segment is punctuation, numbers, or English - use Times New Roman
                run.font.name = "Times New Roman"

    # FOOTER TABLE: Buffer and add side-by-side components
    if left_footer or right_footer:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        cells = table.rows[0].cells
        
        # Helper for table cells
        def add_formatted_runs(cell_p, txt):
            segs = re.split(r'([\u0b80-\u0bff]+)', txt)
            for s in segs:
                if not s: continue
                r = cell_p.add_run(s)
                r.font.size = Pt(12)
                if any('\u0b80' <= char <= '\u0bff' for char in s):
                    r.font.name = "SunTommy"
                    rPr = r.font._element.get_or_add_rPr()
                    rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:cs'), "SunTommy")
                else: r.font.name = "Times New Roman"

        # Left Cell: Date and Place
        for txt, tamil, english in left_footer:
            p = cells[0].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            add_formatted_runs(p, txt)

        # Right Cell: Buffer Signature
        for txt, tamil, english in right_footer:
            p = cells[1].add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = Pt(0)
            add_formatted_runs(p, txt)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
